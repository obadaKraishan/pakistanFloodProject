# content_analysis.py
import pandas as pd
import chardet
from docx import Document
from collections import Counter
from textblob import TextBlob
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
import matplotlib.pyplot as plt
import spacy
from .engagement_analysis import EngagementAnalysis


class ContentAnalysis:
    def __init__(self, processed_dataset_path):
        self.dataset_path = processed_dataset_path
        self.tweets_df = self.load_processed_data()
        # Load SpaCy's English NER model
        self.nlp = spacy.load("en_core_web_sm")

    def load_processed_data(self):
        """Load the processed dataset from an Excel file."""
        try:
            # If your data is in the first sheet, you can omit the sheet_name parameter
            return pd.read_excel(self.dataset_path)
        except Exception as e:
            print(f"Error loading Excel file: {e}")
            return None

    def ngram_analysis(self):
        count_vectorizer = CountVectorizer(ngram_range=(2, 3), max_features=10000)
        X_counts = count_vectorizer.fit_transform(self.tweets_df['cleaned_content'])
        phrases = count_vectorizer.get_feature_names_out()
        return phrases[:10]  # as an example

    def topic_modeling(self):
        """Apply LDA to identify topics and return them."""
        count_vectorizer = CountVectorizer(ngram_range=(1, 2), max_features=10000)
        X_counts = count_vectorizer.fit_transform(self.tweets_df['cleaned_content'])
        lda = LatentDirichletAllocation(n_components=5, random_state=42)
        lda.fit(X_counts)
        feature_names = count_vectorizer.get_feature_names_out()
        topics = []
        for topic_idx, topic in enumerate(lda.components_):
            top_features = [feature_names[i] for i in topic.argsort()[:-10 - 1:-1]]
            topics.append("Topic {}: {}".format(topic_idx + 1, ", ".join(top_features)))
        return topics  # Returns a list of formatted topic strings

    def entity_analysis(self):
        """Extract named entities and analyze their relationships, returning a list of entities."""
        sample_size = min(1000, len(self.tweets_df))  # Adjust sample size dynamically
        entities_list = []

        def extract_entities(text):
            doc = self.nlp(text)
            return [(ent.text, ent.label_) for ent in doc.ents]

        # Apply the function to a dynamic subset of your tweets
        sampled_content = self.tweets_df['cleaned_content'].sample(sample_size).apply(extract_entities)
        for entities in sampled_content:
            entities_list.extend(entities)  # Flatten list of lists into a single list

        return entities_list  # Return a flattened list of entities

    def sentiment_analysis_over_time(self):
        """Analyze sentiment changes over time."""
        # Ensure 'date' column is in datetime format for accurate resampling
        self.tweets_df['date'] = pd.to_datetime(self.tweets_df['date'], format='%Y-%m-%d', errors='coerce')

        # Filter out rows where 'date' is NaT (not a time)
        valid_dates_df = self.tweets_df.dropna(subset=['date'])

        # Now perform resampling on the filtered DataFrame
        sentiment_over_time = valid_dates_df.resample('W', on='date')['sentiment'].mean()

        plt.figure(figsize=(10, 5))
        plt.plot(sentiment_over_time.index, sentiment_over_time, marker='o', linestyle='-', color='b')
        plt.title('Weekly Sentiment Over Time')
        plt.xlabel('Date')
        plt.ylabel('Average Sentiment')
        plt.show()

    def save_analysis_to_doc(self):
        """Save analysis results to a Word document."""
        doc = Document()
        doc.add_heading('Content Analysis Results', 0)

        # Keyword Extraction
        keywords = self.keyword_extraction()
        if keywords is not None:
            doc.add_heading('Keyword Extraction:', level=1)
            doc.add_paragraph(", ".join(keywords))

        # Hashtag Analysis
        hashtags = self.hashtag_analysis()
        if hashtags is not None:
            doc.add_heading('Hashtag Analysis:', level=1)
            for tag, count in hashtags:
                doc.add_paragraph(f"{tag}: {count}")

        # Sentiment Analysis
        sentiments = self.sentiment_analysis()
        doc.add_heading('Sentiment Analysis', level=1)
        for category, count in sentiments.items():
            doc.add_paragraph(f"{category}: {count}")

        # N-Gram Analysis
        phrases = self.ngram_analysis()  # This now returns a list
        doc.add_heading('N-Gram Analysis:', level=1)
        doc.add_paragraph(", ".join(phrases))  # Joining a list works without issues

        # Topic Modeling
        topics = self.topic_modeling()
        doc.add_heading('Topic Modeling', level=1)
        for idx, topic in enumerate(topics, 1):
            doc.add_paragraph(f"Topic {idx}: {', '.join(topic)}")

        # Entity Analysis
        entities = self.entity_analysis()
        doc.add_heading('Entity Analysis', level=1)
        for entity in entities:
            doc.add_paragraph(f"{entity[0]} ({entity[1]})")

        # Crisis Communication Strategy Analysis
        strategies = self.crisis_communication_strategy_analysis()
        if strategies is not None:
            doc.add_heading('Crisis Communication Strategies:', level=1)
            for strategy, count in strategies.items():
                doc.add_paragraph(f"{strategy}: {count}")

        # Save the document
        doc_path = 'content_analysis_results.docx'
        doc.save(doc_path)
        print(f"Analysis results saved to {doc_path}.")

    # Modify your existing methods to return data instead of printing them, for example:
    def keyword_extraction(self):
        """Perform keyword extraction using TF-IDF and return keywords."""
        tfidf_vectorizer = TfidfVectorizer(max_df=0.8, max_features=10000)
        tfidf_matrix = tfidf_vectorizer.fit_transform(self.tweets_df['cleaned_content'])
        keywords = tfidf_vectorizer.get_feature_names_out()
        return keywords[:10]  # Return top 10 keywords as an example

    def hashtag_analysis(self):
        """Analyze hashtags for trending topics and return them."""
        all_hashtags = ','.join(self.tweets_df['hashtags'].dropna()).split(',')
        hashtag_counts = Counter(all_hashtags)
        return hashtag_counts.most_common(10)

    def sentiment_analysis(self):
        """Apply sentiment analysis and return categorized tweets."""
        self.tweets_df['sentiment'] = self.tweets_df['cleaned_content'].apply(lambda tweet: TextBlob(tweet).sentiment.polarity)
        self.tweets_df['sentiment_category'] = self.tweets_df['sentiment'].apply(lambda score: 'Positive' if score > 0 else ('Neutral' if score == 0 else 'Negative'))
        return Counter(self.tweets_df['sentiment_category'])

    def crisis_communication_strategy_analysis(self):
        """
        Analyze tweets for indications of crisis communication strategies based on SCCT
        and return a summary of strategy counts.
        """
        # Define keywords for each SCCT strategy
        strategy_keywords = {
            'denial': ['deny', 'denies', 'not true'],
            'diminishment': ['minor issue', 'not a major problem'],
            'rebuilding': ['apologize', 'sorry', 'compensation'],
            'bolstering': ['support', 'community', 'together']
        }

        # Initialize columns for each strategy in the DataFrame
        for strategy in strategy_keywords:
            self.tweets_df[strategy] = 0

        # Initialize a dictionary to count occurrences of each strategy
        strategy_counts = {strategy: 0 for strategy in strategy_keywords}

        # Iterate through DataFrame to categorize tweets according to the strategies
        for index, row in self.tweets_df.iterrows():
            content = row['cleaned_content'].lower()
            for strategy, keywords in strategy_keywords.items():
                if any(keyword in content for keyword in keywords):
                    self.tweets_df.at[index, strategy] = 1
                    strategy_counts[strategy] += 1  # Increment the count for this strategy

        return strategy_counts

    def perform_all_analyses(self):
        """Run all analysis methods."""
        if self.tweets_df is None:
            print("No data loaded. Exiting analysis.")
            return  # Exit the method if data hasn't been loaded

        print("Starting Content Analysis...")
        self.keyword_extraction()
        self.hashtag_analysis()
        self.sentiment_analysis()
        self.ngram_analysis()
        self.topic_modeling()
        self.sentiment_analysis_over_time()
        self.entity_analysis()
        self.crisis_communication_strategy_analysis()

        # Engagement Analysis
        engagement_analyzer = EngagementAnalysis(self.tweets_df)
        engagement_analyzer.analyze_engagement()
        engagement_analyzer.perform_detailed_engagement_analysis()

        print("Content Analysis Completed.")
